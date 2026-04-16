"""
GAZE Security Platform - Report Generator Service
Word and PDF report generation with report_type and options support.
"""
import os
import tempfile
from datetime import datetime, timezone
from typing import Optional

from jinja2 import Environment, BaseLoader
from flask import current_app

from app.models import Assessment, Finding, Severity


# Severity ordering for consistent sorting
_SEVERITY_ORDER = [
    Severity.CRITICAL,
    Severity.HIGH,
    Severity.MEDIUM,
    Severity.LOW,
    Severity.INFORMATIONAL,
]


def _get_assessment(assessment_id: str) -> Assessment:
    """Fetch assessment or raise 404."""
    return Assessment.query.get_or_404(assessment_id)


def _sort_findings(findings) -> list:
    """Sort findings by severity (critical first)."""
    return sorted(
        findings,
        key=lambda f: _SEVERITY_ORDER.index(f.severity) if f.severity in _SEVERITY_ORDER else 99
    )


def _severity_counts(findings) -> dict:
    """Return {severity_value: count} dict."""
    return {
        sev.value: sum(1 for f in findings if f.severity == sev)
        for sev in _SEVERITY_ORDER
    }


def _report_title(report_type: str) -> str:
    titles = {
        "executive": "Executive Summary Report",
        "technical": "Technical Assessment Report",
        "compliance": "Compliance Assessment Report",
        "full": "Full Security Assessment Report",
        "custom": "Security Assessment Report",
    }
    return titles.get(report_type, "Security Assessment Report")


# =============================================================================
# HTML TEMPLATES
# =============================================================================

_BASE_CSS = """
body { font-family: Arial, sans-serif; margin: 40px; color: #333; line-height: 1.6; }
h1 { color: #1a1a2e; border-bottom: 2px solid #1a1a2e; padding-bottom: 10px; }
h2 { color: #16213e; margin-top: 30px; }
h3 { color: #0f3460; }
h4 { color: #333; margin: 10px 0 5px; }
table { border-collapse: collapse; width: 100%; margin: 15px 0; }
th, td { border: 1px solid #ddd; padding: 10px; text-align: left; }
th { background-color: #1a1a2e; color: white; }
.severity-critical { background-color: #dc2626; color: white; font-weight: bold; padding: 2px 6px; border-radius: 3px; }
.severity-high { background-color: #ea580c; color: white; font-weight: bold; padding: 2px 6px; border-radius: 3px; }
.severity-medium { background-color: #ca8a04; color: white; font-weight: bold; padding: 2px 6px; border-radius: 3px; }
.severity-low { background-color: #16a34a; color: white; font-weight: bold; padding: 2px 6px; border-radius: 3px; }
.severity-informational { background-color: #2563eb; color: white; padding: 2px 6px; border-radius: 3px; }
.finding { margin: 20px 0; padding: 15px; border: 1px solid #ddd; border-radius: 5px; page-break-inside: avoid; }
.finding-title { font-size: 1.1em; font-weight: bold; margin-bottom: 10px; }
.meta { color: #555; font-size: 0.9em; background: #f8f8f8; padding: 12px; border-radius: 4px; margin-bottom: 20px; }
.meta p { margin: 4px 0; }
.summary-box { display: inline-block; padding: 10px 20px; margin: 5px; border-radius: 6px; text-align: center; min-width: 80px; }
.summary-box .count { font-size: 2em; font-weight: bold; }
.summary-box .label { font-size: 0.8em; }
.exec-highlight { background: #f0f4ff; border-left: 4px solid #1a1a2e; padding: 12px 16px; margin: 15px 0; }
@page { margin: 2cm; }
"""

_FULL_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{{ title }}</title>
    <style>{{ css }}</style>
</head>
<body>
    <h1>{{ title }}</h1>

    <div class="meta">
        <p><strong>Assessment:</strong> {{ assessment.name }}</p>
        <p><strong>Asset:</strong> {{ assessment.asset.name if assessment.asset else 'N/A' }}</p>
        <p><strong>Type:</strong> {{ assessment.assessment_type.value.replace('_', ' ').title() }}</p>
        <p><strong>Status:</strong> {{ assessment.status.value.replace('_', ' ').title() }}</p>
        <p><strong>Generated:</strong> {{ now }}</p>
    </div>

    {% if assessment.executive_summary and include_executive %}
    <h2>Executive Summary</h2>
    <div class="exec-highlight">{{ assessment.executive_summary }}</div>
    {% endif %}

    <h2>Findings Summary</h2>
    <table>
        <tr>
            <th>Severity</th>
            <th>Count</th>
        </tr>
        {% for sev, count in severity_counts.items() %}
        <tr>
            <td><span class="severity-{{ sev }}">{{ sev.upper() }}</span></td>
            <td>{{ count }}</td>
        </tr>
        {% endfor %}
    </table>

    {% if include_metrics %}
    <h2>Risk Metrics</h2>
    <p>Total open findings: <strong>{{ findings | length }}</strong></p>
    {% endif %}

    {% if include_findings %}
    <h2>Detailed Findings</h2>
    {% for finding in findings %}
    <div class="finding">
        <div class="finding-title">{{ loop.index }}. {{ finding.title }}</div>
        <table>
            <tr>
                <td><strong>Severity</strong></td>
                <td><span class="severity-{{ finding.severity.value }}">{{ finding.severity.value.upper() }}</span></td>
            </tr>
            <tr><td><strong>Status</strong></td><td>{{ finding.status.value.replace('_', ' ').title() }}</td></tr>
            {% if finding.cvss_score %}<tr><td><strong>CVSS</strong></td><td>{{ finding.cvss_score }}</td></tr>{% endif %}
            {% if finding.cwe_id %}<tr><td><strong>CWE</strong></td><td>{{ finding.cwe_id }}</td></tr>{% endif %}
            {% if finding.cve_id %}<tr><td><strong>CVE</strong></td><td>{{ finding.cve_id }}</td></tr>{% endif %}
            {% if finding.owasp_category %}<tr><td><strong>OWASP</strong></td><td>{{ finding.owasp_category }}</td></tr>{% endif %}
            {% if finding.affected_component %}<tr><td><strong>Component</strong></td><td>{{ finding.affected_component }}</td></tr>{% endif %}
        </table>

        {% if finding.description %}
        <h4>Description</h4>
        <p>{{ finding.description }}</p>
        {% endif %}

        {% if finding.impact %}
        <h4>Impact</h4>
        <p>{{ finding.impact }}</p>
        {% endif %}

        {% if include_remediation and finding.recommendation %}
        <h4>Recommendation</h4>
        <p>{{ finding.recommendation }}</p>
        {% endif %}

        {% if include_evidence and finding.evidence %}
        <h4>Evidence</h4>
        <p>{{ finding.evidence }}</p>
        {% endif %}
    </div>
    {% endfor %}
    {% endif %}

    {% if include_timeline %}
    <h2>Remediation Timeline</h2>
    <table>
        <tr>
            <th>Finding</th>
            <th>Severity</th>
            <th>Due Date</th>
            <th>Status</th>
        </tr>
        {% for finding in findings %}
        <tr>
            <td>{{ finding.title }}</td>
            <td><span class="severity-{{ finding.severity.value }}">{{ finding.severity.value.upper() }}</span></td>
            <td>{{ finding.due_date.strftime('%Y-%m-%d') if finding.due_date else 'TBD' }}</td>
            <td>{{ finding.status.value.replace('_', ' ').title() }}</td>
        </tr>
        {% endfor %}
    </table>
    {% endif %}

</body>
</html>
"""

_EXECUTIVE_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{{ title }}</title>
    <style>{{ css }}</style>
</head>
<body>
    <h1>{{ title }}</h1>

    <div class="meta">
        <p><strong>Assessment:</strong> {{ assessment.name }}</p>
        <p><strong>Asset:</strong> {{ assessment.asset.name if assessment.asset else 'N/A' }}</p>
        <p><strong>Generated:</strong> {{ now }}</p>
    </div>

    {% if assessment.executive_summary %}
    <h2>Executive Summary</h2>
    <div class="exec-highlight">{{ assessment.executive_summary }}</div>
    {% endif %}

    <h2>Risk Overview</h2>
    <table>
        <tr>
            <th>Severity</th>
            <th>Count</th>
            <th>Risk Level</th>
        </tr>
        {% for sev, count in severity_counts.items() %}
        <tr>
            <td><span class="severity-{{ sev }}">{{ sev.upper() }}</span></td>
            <td>{{ count }}</td>
            <td>{% if sev in ['critical','high'] %}Immediate Action Required{% elif sev == 'medium' %}Action Required{% else %}Monitor{% endif %}</td>
        </tr>
        {% endfor %}
    </table>

    <h2>Key Findings</h2>
    {% for finding in findings[:5] %}
    <div class="finding">
        <div class="finding-title">
            <span class="severity-{{ finding.severity.value }}">{{ finding.severity.value.upper() }}</span>
            &nbsp;{{ finding.title }}
        </div>
        {% if finding.impact %}<p><strong>Impact:</strong> {{ finding.impact }}</p>{% endif %}
        {% if include_remediation and finding.recommendation %}
        <p><strong>Recommendation:</strong> {{ finding.recommendation }}</p>
        {% endif %}
    </div>
    {% endfor %}
    {% if findings | length > 5 %}
    <p><em>... and {{ (findings | length) - 5 }} additional findings. See the technical report for full details.</em></p>
    {% endif %}
</body>
</html>
"""


def _build_template_context(assessment: Assessment, report_type: str, options: dict) -> dict:
    """Build Jinja2 template context from assessment + options."""
    findings = _sort_findings(assessment.findings)
    counts = _severity_counts(assessment.findings)

    # Map option keys coming from frontend
    include_executive = options.get("include_executive", True)
    include_findings = report_type in ("technical", "full", "compliance", "custom")
    include_remediation = options.get("include_remediation", True)
    include_evidence = options.get("include_evidence", False)
    include_metrics = options.get("include_metrics", True)
    include_timeline = options.get("include_timeline", False)

    return {
        "assessment": assessment,
        "findings": findings,
        "severity_counts": counts,
        "title": _report_title(report_type),
        "now": datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"),
        "css": _BASE_CSS,
        "include_executive": include_executive,
        "include_findings": include_findings,
        "include_remediation": include_remediation,
        "include_evidence": include_evidence,
        "include_metrics": include_metrics,
        "include_timeline": include_timeline,
    }


def _render_html(assessment: Assessment, report_type: str, options: dict) -> str:
    """Render the appropriate HTML template."""
    env = Environment(loader=BaseLoader(), autoescape=True)
    ctx = _build_template_context(assessment, report_type, options)

    # Executive reports get a condensed template
    raw = _EXECUTIVE_TEMPLATE if report_type == "executive" else _FULL_TEMPLATE
    tmpl = env.from_string(raw)
    return tmpl.render(**ctx)


# =============================================================================
# PUBLIC API
# =============================================================================

class ReportGenerator:
    """Service for generating assessment reports in multiple formats."""

    @staticmethod
    def generate_word_report(
        assessment_id: str,
        report_type: str = "technical",
        options: Optional[dict] = None,
    ) -> str:
        """
        Generate a Word (.docx) report.

        Args:
            assessment_id: UUID string of the assessment.
            report_type: One of executive | technical | compliance | full | custom.
            options: Dict of inclusion flags (include_remediation, include_evidence, etc.)

        Returns:
            Absolute path to the generated .docx temp file.
        """
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        options = options or {}
        assessment = _get_assessment(assessment_id)

        include_remediation = options.get("include_remediation", True)
        include_evidence = options.get("include_evidence", False)
        include_timeline = options.get("include_timeline", False)
        include_findings = report_type in ("technical", "full", "compliance", "custom")

        doc = Document()

        # Title
        title = doc.add_heading(_report_title(report_type), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Assessment metadata
        doc.add_heading("Assessment Details", level=1)
        doc.add_paragraph(f"Assessment: {assessment.name}")
        doc.add_paragraph(f"Asset: {assessment.asset.name if assessment.asset else 'N/A'}")
        doc.add_paragraph(f"Type: {assessment.assessment_type.value.replace('_', ' ').title()}")
        doc.add_paragraph(f"Status: {assessment.status.value.replace('_', ' ').title()}")
        doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")

        # Executive summary
        if assessment.executive_summary:
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(assessment.executive_summary)

        # Findings summary table
        doc.add_heading("Findings Summary", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Severity"
        hdr[1].text = "Count"

        for sev in _SEVERITY_ORDER:
            count = sum(1 for f in assessment.findings if f.severity == sev)
            row = table.add_row().cells
            row[0].text = sev.value.upper()
            row[1].text = str(count)

        # Detailed findings
        if include_findings:
            doc.add_heading("Detailed Findings", level=1)
            sorted_findings = _sort_findings(assessment.findings)

            for sev in _SEVERITY_ORDER:
                group = [f for f in sorted_findings if f.severity == sev]
                if not group:
                    continue
                doc.add_heading(f"{sev.value.upper()} Severity Findings", level=2)

                for i, finding in enumerate(group, 1):
                    doc.add_heading(f"{i}. {finding.title}", level=3)

                    details_table = doc.add_table(rows=0, cols=2)
                    details_table.style = "Table Grid"

                    details = [
                        ("Severity", finding.severity.value.upper()),
                        ("Status", finding.status.value.replace("_", " ").title()),
                        ("CVSS Score", str(finding.cvss_score) if finding.cvss_score else "N/A"),
                        ("CWE", finding.cwe_id or "N/A"),
                        ("CVE", finding.cve_id or "N/A"),
                        ("OWASP", finding.owasp_category or "N/A"),
                        ("Affected Component", finding.affected_component or "N/A"),
                    ]
                    for label, value in details:
                        row = details_table.add_row().cells
                        row[0].text = label
                        row[1].text = value

                    doc.add_paragraph()

                    if finding.description:
                        doc.add_heading("Description", level=4)
                        doc.add_paragraph(finding.description)
                    if finding.impact:
                        doc.add_heading("Impact", level=4)
                        doc.add_paragraph(finding.impact)
                    if include_remediation and finding.recommendation:
                        doc.add_heading("Recommendation", level=4)
                        doc.add_paragraph(finding.recommendation)
                    if include_evidence and getattr(finding, "evidence", None):
                        doc.add_heading("Evidence", level=4)
                        doc.add_paragraph(finding.evidence)

                    doc.add_paragraph()

        # Remediation timeline
        if include_timeline:
            doc.add_heading("Remediation Timeline", level=1)
            tl_table = doc.add_table(rows=1, cols=4)
            tl_table.style = "Table Grid"
            hdr = tl_table.rows[0].cells
            hdr[0].text = "Finding"
            hdr[1].text = "Severity"
            hdr[2].text = "Due Date"
            hdr[3].text = "Status"
            for finding in _sort_findings(assessment.findings):
                row = tl_table.add_row().cells
                row[0].text = finding.title
                row[1].text = finding.severity.value.upper()
                row[2].text = (
                    finding.due_date.strftime("%Y-%m-%d")
                    if getattr(finding, "due_date", None)
                    else "TBD"
                )
                row[3].text = finding.status.value.replace("_", " ").title()

        fd, filepath = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc.save(filepath)
        return filepath

    @staticmethod
    def generate_pdf_report(
        assessment_id: str,
        report_type: str = "technical",
        options: Optional[dict] = None,
    ) -> str:
        """
        Generate a PDF report via the sandbox service (WeasyPrint fallback).

        Args:
            assessment_id: UUID string of the assessment.
            report_type: One of executive | technical | compliance | full | custom.
            options: Dict of inclusion flags.

        Returns:
            Absolute path to the generated .pdf temp file.
        """
        import requests as req

        options = options or {}
        assessment = _get_assessment(assessment_id)
        html_content = _render_html(assessment, report_type, options)

        html_path = "/tmp/report_input.html"
        fd, pdf_path = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)

        with open(html_path, "w") as f:
            f.write(html_content)

        sandbox_host = current_app.config.get("PDF_SANDBOX_HOST", "pdf-sandbox")
        sandbox_port = current_app.config.get("PDF_SANDBOX_PORT", 8001)

        try:
            resp = req.post(
                f"http://{sandbox_host}:{sandbox_port}/generate",
                json={"html_path": html_path, "output_path": pdf_path},
                timeout=30,
            )
            resp.raise_for_status()
        except Exception:
            # Fallback for dev environments
            from weasyprint import HTML
            HTML(string=html_content).write_pdf(pdf_path)

        return pdf_path

    @staticmethod
    def generate_html_report(
        assessment_id: str,
        report_type: str = "technical",
        options: Optional[dict] = None,
    ) -> str:
        """
        Write the HTML report to a temp file and return the path.

        Args:
            assessment_id: UUID string of the assessment.
            report_type: One of executive | technical | compliance | full | custom.
            options: Dict of inclusion flags.

        Returns:
            Absolute path to the generated .html temp file.
        """
        options = options or {}
        assessment = _get_assessment(assessment_id)
        html_content = _render_html(assessment, report_type, options)

        fd, filepath = tempfile.mkstemp(suffix=".html")
        with os.fdopen(fd, "w") as f:
            f.write(html_content)

        return filepath

    @staticmethod
    def generate_html_preview(
        assessment_id: str,
        report_type: str = "technical",
        options: Optional[dict] = None,
    ) -> str:
        """
        Return HTML content as a string (for in-browser viewing, no file written).

        Args:
            assessment_id: UUID string of the assessment.
            report_type: One of executive | technical | compliance | full | custom.
            options: Dict of inclusion flags.

        Returns:
            HTML string.
        """
        options = options or {}
        assessment = _get_assessment(assessment_id)
        return _render_html(assessment, report_type, options)

    # ------------------------------------------------------------------
    # Keep old 1-arg signatures working so legacy export routes don't break
    # ------------------------------------------------------------------

    @staticmethod
    def _generate_html_report(assessment: Assessment) -> str:
        """Legacy internal method — kept for backward compat."""
        return _render_html(assessment, "technical", {})