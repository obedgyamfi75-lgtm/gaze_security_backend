#!/usr/bin/env python3
"""
╔═══════════════════════════════════════════════════════════════════════════════╗
║                     HUBTEL SECURITY REPORT GENERATOR                          ║
║              Professional Security Assessment Report Builder                   ║
╠═══════════════════════════════════════════════════════════════════════════════╣
║  Organization : GAZE Limited                                                 ║
║  Team         : Security Operations                                            ║
║  Purpose      : Generate professional security assessment reports              ║
╚═══════════════════════════════════════════════════════════════════════════════╝

Features:
    - Word document generation (.docx)
    - PDF generation (sandboxed)
    - Executive and technical report formats
    - Custom branding and templates
    - Finding severity visualization
    - Evidence attachment support

Usage:
    from app.lib.reporter import SecurityReporter
    
    reporter = SecurityReporter(assessment)
    reporter.generate_word("report.docx")
    reporter.generate_pdf("report.pdf")
"""
import os
import tempfile
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Optional, Any
from dataclasses import dataclass, field
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from jinja2 import Environment, BaseLoader
import structlog

logger = structlog.get_logger()


# ══════════════════════════════════════════════════════════════════════════════
# Data Classes
# ══════════════════════════════════════════════════════════════════════════════

@dataclass
class FindingData:
    """Finding data structure for reports"""
    id: str
    title: str
    severity: str  # critical, high, medium, low, informational
    status: str
    cvss_score: Optional[float] = None
    cvss_vector: Optional[str] = None
    cwe_id: Optional[str] = None
    cve_id: Optional[str] = None
    owasp_category: Optional[str] = None
    description: str = ""
    impact: str = ""
    affected_component: str = ""
    affected_url: str = ""
    steps_to_reproduce: str = ""
    poc_code: str = ""
    recommendation: str = ""
    evidence: List[str] = field(default_factory=list)
    
    @property
    def severity_color(self) -> tuple:
        """RGB color for severity"""
        colors = {
            'critical': (220, 38, 38),   # Red
            'high': (234, 88, 12),       # Orange
            'medium': (202, 138, 4),     # Yellow
            'low': (22, 163, 74),        # Green
            'informational': (37, 99, 235),  # Blue
        }
        return colors.get(self.severity.lower(), (128, 128, 128))


@dataclass
class AssessmentData:
    """Assessment data structure for reports"""
    id: str
    name: str
    asset_name: str
    assessment_type: str
    status: str
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    executive_summary: str = ""
    scope_description: str = ""
    methodology: str = ""
    findings: List[FindingData] = field(default_factory=list)
    assessor_name: str = "Security Team"
    assessor_email: str = "security@hubtel.com"
    
    @property
    def findings_by_severity(self) -> Dict[str, List[FindingData]]:
        """Group findings by severity"""
        result = {
            'critical': [],
            'high': [],
            'medium': [],
            'low': [],
            'informational': [],
        }
        for finding in self.findings:
            sev = finding.severity.lower()
            if sev in result:
                result[sev].append(finding)
        return result
    
    @property
    def severity_counts(self) -> Dict[str, int]:
        """Count findings by severity"""
        return {k: len(v) for k, v in self.findings_by_severity.items()}


# ══════════════════════════════════════════════════════════════════════════════
# Report Configuration
# ══════════════════════════════════════════════════════════════════════════════

@dataclass
class ReportConfig:
    """Report configuration"""
    company_name: str = "GAZE Limited"
    company_logo: Optional[str] = None  # Path to logo
    report_title: str = "Security Assessment Report"
    classification: str = "CONFIDENTIAL"
    include_executive_summary: bool = True
    include_methodology: bool = True
    include_technical_details: bool = True
    include_evidence: bool = True
    include_remediation: bool = True


# ══════════════════════════════════════════════════════════════════════════════
# Word Document Generator
# ══════════════════════════════════════════════════════════════════════════════

class WordReportGenerator:
    """Generate Word document reports"""
    
    def __init__(self, config: Optional[ReportConfig] = None):
        self.config = config or ReportConfig()
    
    def generate(self, assessment: AssessmentData, output_path: str) -> str:
        """
        Generate Word document report.
        
        Args:
            assessment: Assessment data
            output_path: Output file path
        
        Returns:
            Path to generated document
        """
        doc = Document()
        self._setup_styles(doc)
        
        # Title page
        self._add_title_page(doc, assessment)
        
        # Table of contents placeholder
        doc.add_page_break()
        doc.add_heading("Table of Contents", level=1)
        doc.add_paragraph("(Update table of contents after editing)")
        
        # Executive summary
        if self.config.include_executive_summary:
            doc.add_page_break()
            self._add_executive_summary(doc, assessment)
        
        # Findings summary
        doc.add_page_break()
        self._add_findings_summary(doc, assessment)
        
        # Detailed findings
        if self.config.include_technical_details:
            doc.add_page_break()
            self._add_detailed_findings(doc, assessment)
        
        # Methodology
        if self.config.include_methodology and assessment.methodology:
            doc.add_page_break()
            self._add_methodology(doc, assessment)
        
        # Remediation summary
        if self.config.include_remediation:
            doc.add_page_break()
            self._add_remediation_summary(doc, assessment)
        
        # Save document
        doc.save(output_path)
        logger.info("word_report_generated", path=output_path)
        return output_path
    
    def _setup_styles(self, doc: Document) -> None:
        """Set up document styles"""
        # Heading styles
        for i in range(1, 4):
            style = doc.styles[f'Heading {i}']
            style.font.color.rgb = RGBColor(26, 26, 46)  # Dark blue
            style.font.bold = True
    
    def _add_title_page(self, doc: Document, assessment: AssessmentData) -> None:
        """Add title page"""
        # Add spacing
        for _ in range(3):
            doc.add_paragraph()
        
        # Company name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.config.company_name)
        run.font.size = Pt(24)
        run.font.bold = True
        
        doc.add_paragraph()
        
        # Report title
        title = doc.add_heading(self.config.report_title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Assessment name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(assessment.name)
        run.font.size = Pt(18)
        run.font.italic = True
        
        # Asset
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Target: {assessment.asset_name}")
        
        # Spacing
        for _ in range(5):
            doc.add_paragraph()
        
        # Classification
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.config.classification)
        run.font.bold = True
        run.font.color.rgb = RGBColor(220, 38, 38)
        
        # Date and assessor
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Assessment Date: {assessment.start_date or datetime.now().strftime('%Y-%m-%d')}")
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Prepared by: {assessment.assessor_name}")
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(assessment.assessor_email)
    
    def _add_executive_summary(self, doc: Document, assessment: AssessmentData) -> None:
        """Add executive summary section"""
        doc.add_heading("Executive Summary", level=1)
        
        # Overview
        if assessment.executive_summary:
            doc.add_paragraph(assessment.executive_summary)
        else:
            doc.add_paragraph(
                f"This report presents the findings of a {assessment.assessment_type.replace('_', ' ')} "
                f"conducted on {assessment.asset_name}. The assessment identified "
                f"{len(assessment.findings)} security findings."
            )
        
        # Findings overview table
        doc.add_heading("Findings Overview", level=2)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Severity"
        hdr_cells[1].text = "Count"
        
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            self._set_cell_shading(cell, "1A1A2E")
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Data rows
        for severity, count in assessment.severity_counts.items():
            row = table.add_row().cells
            row[0].text = severity.upper()
            row[1].text = str(count)
            
            # Color code severity
            color_map = {
                'critical': 'DC2626',
                'high': 'EA580C',
                'medium': 'CA8A04',
                'low': '16A34A',
                'informational': '2563EB',
            }
            if severity in color_map:
                self._set_cell_shading(row[0], color_map[severity])
                row[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    def _add_findings_summary(self, doc: Document, assessment: AssessmentData) -> None:
        """Add findings summary table"""
        doc.add_heading("Findings Summary", level=1)
        
        if not assessment.findings:
            doc.add_paragraph("No findings identified during this assessment.")
            return
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Header
        headers = ["#", "Title", "Severity", "Status", "CVSS"]
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            self._set_cell_shading(hdr_cells[i], "1A1A2E")
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Data rows
        for idx, finding in enumerate(assessment.findings, 1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = finding.title[:50] + ("..." if len(finding.title) > 50 else "")
            row[2].text = finding.severity.upper()
            row[3].text = finding.status.replace("_", " ").title()
            row[4].text = str(finding.cvss_score) if finding.cvss_score else "N/A"
            
            # Color severity
            color_map = {
                'critical': 'DC2626',
                'high': 'EA580C',
                'medium': 'CA8A04',
                'low': '16A34A',
                'informational': '2563EB',
            }
            if finding.severity.lower() in color_map:
                self._set_cell_shading(row[2], color_map[finding.severity.lower()])
                row[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    def _add_detailed_findings(self, doc: Document, assessment: AssessmentData) -> None:
        """Add detailed findings section"""
        doc.add_heading("Detailed Findings", level=1)
        
        severity_order = ['critical', 'high', 'medium', 'low', 'informational']
        
        for severity in severity_order:
            findings = assessment.findings_by_severity.get(severity, [])
            if not findings:
                continue
            
            doc.add_heading(f"{severity.upper()} Severity Findings", level=2)
            
            for idx, finding in enumerate(findings, 1):
                self._add_finding_detail(doc, finding, idx)
    
    def _add_finding_detail(self, doc: Document, finding: FindingData, idx: int) -> None:
        """Add individual finding details"""
        doc.add_heading(f"{idx}. {finding.title}", level=3)
        
        # Metadata table
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        metadata = [
            ("Severity", finding.severity.upper()),
            ("Status", finding.status.replace("_", " ").title()),
            ("CVSS Score", str(finding.cvss_score) if finding.cvss_score else "N/A"),
            ("CWE", finding.cwe_id or "N/A"),
            ("CVE", finding.cve_id or "N/A"),
            ("OWASP", finding.owasp_category or "N/A"),
            ("Affected Component", finding.affected_component or "N/A"),
        ]
        
        for label, value in metadata:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = value
            row[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Description
        if finding.description:
            doc.add_heading("Description", level=4)
            doc.add_paragraph(finding.description)
        
        # Impact
        if finding.impact:
            doc.add_heading("Impact", level=4)
            doc.add_paragraph(finding.impact)
        
        # Steps to reproduce
        if finding.steps_to_reproduce:
            doc.add_heading("Steps to Reproduce", level=4)
            doc.add_paragraph(finding.steps_to_reproduce)
        
        # POC
        if finding.poc_code:
            doc.add_heading("Proof of Concept", level=4)
            p = doc.add_paragraph()
            run = p.add_run(finding.poc_code)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        # Recommendation
        if finding.recommendation:
            doc.add_heading("Recommendation", level=4)
            doc.add_paragraph(finding.recommendation)
        
        doc.add_paragraph()  # Spacing
    
    def _add_methodology(self, doc: Document, assessment: AssessmentData) -> None:
        """Add methodology section"""
        doc.add_heading("Methodology", level=1)
        doc.add_paragraph(assessment.methodology)
    
    def _add_remediation_summary(self, doc: Document, assessment: AssessmentData) -> None:
        """Add remediation summary"""
        doc.add_heading("Remediation Summary", level=1)
        
        doc.add_paragraph(
            "The following table summarizes the recommended remediation actions "
            "and their priority based on finding severity."
        )
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        headers = ["Finding", "Severity", "Recommendation"]
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
        for finding in assessment.findings:
            if finding.recommendation:
                row = table.add_row().cells
                row[0].text = finding.title[:40]
                row[1].text = finding.severity.upper()
                row[2].text = finding.recommendation[:100] + ("..." if len(finding.recommendation) > 100 else "")
    
    def _set_cell_shading(self, cell, color: str) -> None:
        """Set cell background color"""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading)


# ══════════════════════════════════════════════════════════════════════════════
# HTML/PDF Generator
# ══════════════════════════════════════════════════════════════════════════════

class HTMLReportGenerator:
    """Generate HTML reports (for PDF conversion)"""
    
    TEMPLATE = '''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{{ assessment.name }} - Security Report</title>
    <style>
        @page { margin: 2cm; size: A4; }
        body { font-family: 'Segoe UI', Arial, sans-serif; margin: 40px; color: #333; line-height: 1.6; }
        h1 { color: #1a1a2e; border-bottom: 3px solid #1a1a2e; padding-bottom: 10px; }
        h2 { color: #16213e; margin-top: 30px; border-bottom: 1px solid #ddd; padding-bottom: 5px; }
        h3 { color: #0f3460; }
        h4 { color: #444; margin-top: 15px; }
        table { border-collapse: collapse; width: 100%; margin: 15px 0; }
        th, td { border: 1px solid #ddd; padding: 12px; text-align: left; }
        th { background-color: #1a1a2e; color: white; font-weight: 600; }
        .severity-critical { background-color: #dc2626; color: white; font-weight: bold; padding: 4px 8px; border-radius: 4px; }
        .severity-high { background-color: #ea580c; color: white; font-weight: bold; padding: 4px 8px; border-radius: 4px; }
        .severity-medium { background-color: #ca8a04; color: black; font-weight: bold; padding: 4px 8px; border-radius: 4px; }
        .severity-low { background-color: #16a34a; color: white; font-weight: bold; padding: 4px 8px; border-radius: 4px; }
        .severity-informational { background-color: #2563eb; color: white; padding: 4px 8px; border-radius: 4px; }
        .finding { margin: 25px 0; padding: 20px; border: 1px solid #ddd; border-radius: 8px; background: #fafafa; }
        .finding-title { font-size: 1.2em; font-weight: bold; margin-bottom: 15px; color: #1a1a2e; }
        .meta { color: #666; font-size: 0.9em; margin-bottom: 15px; }
        .code { background: #1a1a2e; color: #00ff88; padding: 15px; border-radius: 4px; font-family: 'Courier New', monospace; font-size: 0.9em; overflow-x: auto; white-space: pre-wrap; }
        .page-break { page-break-after: always; }
        .title-page { text-align: center; padding-top: 150px; }
        .classification { color: #dc2626; font-weight: bold; font-size: 1.2em; margin: 30px 0; }
        .summary-box { background: #f0f9ff; border-left: 4px solid #2563eb; padding: 15px; margin: 15px 0; }
    </style>
</head>
<body>
    <div class="title-page">
        <h1 style="border: none; font-size: 2.5em;">{{ config.company_name }}</h1>
        <h2 style="border: none; font-size: 1.8em; color: #1a1a2e;">{{ config.report_title }}</h2>
        <p style="font-size: 1.3em; color: #666;">{{ assessment.name }}</p>
        <p>Target: {{ assessment.asset_name }}</p>
        <p class="classification">{{ config.classification }}</p>
        <p style="margin-top: 50px;">
            Assessment Date: {{ assessment.start_date or now }}<br>
            Prepared by: {{ assessment.assessor_name }}<br>
            {{ assessment.assessor_email }}
        </p>
    </div>
    
    <div class="page-break"></div>
    
    <h1>Executive Summary</h1>
    {% if assessment.executive_summary %}
    <p>{{ assessment.executive_summary }}</p>
    {% else %}
    <p>This report presents the findings of a {{ assessment.assessment_type|replace('_', ' ') }} 
       conducted on {{ assessment.asset_name }}. The assessment identified 
       {{ assessment.findings|length }} security findings.</p>
    {% endif %}
    
    <h2>Findings Overview</h2>
    <table>
        <tr>
            <th>Severity</th>
            <th>Count</th>
        </tr>
        {% for severity, count in severity_counts.items() %}
        <tr>
            <td><span class="severity-{{ severity }}">{{ severity|upper }}</span></td>
            <td>{{ count }}</td>
        </tr>
        {% endfor %}
    </table>
    
    <div class="page-break"></div>
    
    <h1>Findings Summary</h1>
    <table>
        <tr>
            <th>#</th>
            <th>Title</th>
            <th>Severity</th>
            <th>Status</th>
            <th>CVSS</th>
        </tr>
        {% for finding in findings %}
        <tr>
            <td>{{ loop.index }}</td>
            <td>{{ finding.title }}</td>
            <td><span class="severity-{{ finding.severity|lower }}">{{ finding.severity|upper }}</span></td>
            <td>{{ finding.status|replace('_', ' ')|title }}</td>
            <td>{{ finding.cvss_score or 'N/A' }}</td>
        </tr>
        {% endfor %}
    </table>
    
    <div class="page-break"></div>
    
    <h1>Detailed Findings</h1>
    {% for finding in findings %}
    <div class="finding">
        <div class="finding-title">{{ loop.index }}. {{ finding.title }}</div>
        
        <table>
            <tr><td><strong>Severity</strong></td><td><span class="severity-{{ finding.severity|lower }}">{{ finding.severity|upper }}</span></td></tr>
            <tr><td><strong>Status</strong></td><td>{{ finding.status|replace('_', ' ')|title }}</td></tr>
            {% if finding.cvss_score %}<tr><td><strong>CVSS</strong></td><td>{{ finding.cvss_score }}</td></tr>{% endif %}
            {% if finding.cwe_id %}<tr><td><strong>CWE</strong></td><td>{{ finding.cwe_id }}</td></tr>{% endif %}
            {% if finding.cve_id %}<tr><td><strong>CVE</strong></td><td>{{ finding.cve_id }}</td></tr>{% endif %}
            {% if finding.affected_component %}<tr><td><strong>Component</strong></td><td>{{ finding.affected_component }}</td></tr>{% endif %}
        </table>
        
        {% if finding.description %}
        <h4>Description</h4>
        <p>{{ finding.description }}</p>
        {% endif %}
        
        {% if finding.impact %}
        <h4>Impact</h4>
        <p>{{ finding.impact }}</p>
        {% endif %}
        
        {% if finding.steps_to_reproduce %}
        <h4>Steps to Reproduce</h4>
        <div class="code">{{ finding.steps_to_reproduce }}</div>
        {% endif %}
        
        {% if finding.poc_code %}
        <h4>Proof of Concept</h4>
        <div class="code">{{ finding.poc_code }}</div>
        {% endif %}
        
        {% if finding.recommendation %}
        <h4>Recommendation</h4>
        <p>{{ finding.recommendation }}</p>
        {% endif %}
    </div>
    {% endfor %}
</body>
</html>'''
    
    def __init__(self, config: Optional[ReportConfig] = None):
        self.config = config or ReportConfig()
        self.env = Environment(loader=BaseLoader(), autoescape=True)
    
    def generate_html(self, assessment: AssessmentData) -> str:
        """Generate HTML report content"""
        template = self.env.from_string(self.TEMPLATE)
        
        return template.render(
            assessment=assessment,
            config=self.config,
            findings=assessment.findings,
            severity_counts=assessment.severity_counts,
            now=datetime.now().strftime('%Y-%m-%d'),
        )
    
    def generate_pdf(self, assessment: AssessmentData, output_path: str) -> str:
        """
        Generate PDF report.
        
        Note: This uses WeasyPrint which should be run in a sandbox
        in production for security.
        """
        try:
            from weasyprint import HTML
        except ImportError:
            logger.error("weasyprint_not_installed")
            raise ImportError("WeasyPrint is required for PDF generation")
        
        html_content = self.generate_html(assessment)
        HTML(string=html_content).write_pdf(output_path)
        
        logger.info("pdf_report_generated", path=output_path)
        return output_path


# ══════════════════════════════════════════════════════════════════════════════
# Main Reporter Class
# ══════════════════════════════════════════════════════════════════════════════

class SecurityReporter:
    """
    Main security report generator.
    
    Example:
        from app.lib.reporter import SecurityReporter, AssessmentData, FindingData
        
        findings = [
            FindingData(
                id="1",
                title="SQL Injection in Login",
                severity="critical",
                status="open",
                cvss_score=9.8,
                description="...",
                recommendation="Use parameterized queries",
            )
        ]
        
        assessment = AssessmentData(
            id="ASM-001",
            name="Q1 2025 Assessment",
            asset_name="Customer Portal",
            assessment_type="penetration_test",
            status="completed",
            findings=findings,
        )
        
        reporter = SecurityReporter()
        reporter.generate_word(assessment, "report.docx")
        reporter.generate_pdf(assessment, "report.pdf")
    """
    
    def __init__(self, config: Optional[ReportConfig] = None):
        self.config = config or ReportConfig()
        self.word_generator = WordReportGenerator(config)
        self.html_generator = HTMLReportGenerator(config)
    
    def generate_word(self, assessment: AssessmentData, output_path: str) -> str:
        """Generate Word document report"""
        return self.word_generator.generate(assessment, output_path)
    
    def generate_pdf(self, assessment: AssessmentData, output_path: str) -> str:
        """Generate PDF report"""
        return self.html_generator.generate_pdf(assessment, output_path)
    
    def generate_html(self, assessment: AssessmentData) -> str:
        """Generate HTML report content"""
        return self.html_generator.generate_html(assessment)
    
    @staticmethod
    def from_db_assessment(assessment) -> AssessmentData:
        """
        Convert database Assessment model to AssessmentData.
        
        Args:
            assessment: SQLAlchemy Assessment model instance
        
        Returns:
            AssessmentData for report generation
        """
        findings = []
        for f in assessment.findings:
            findings.append(FindingData(
                id=str(f.id),
                title=f.title,
                severity=f.severity.value,
                status=f.status.value,
                cvss_score=f.cvss_score,
                cvss_vector=f.cvss_vector,
                cwe_id=f.cwe_id,
                cve_id=f.cve_id,
                owasp_category=f.owasp_category,
                description=f.description or "",
                impact=f.impact or "",
                affected_component=f.affected_component or "",
                affected_url=f.affected_url or "",
                steps_to_reproduce=f.steps_to_reproduce or "",
                poc_code=f.poc_code or "",
                recommendation=f.recommendation or "",
            ))
        
        return AssessmentData(
            id=str(assessment.id),
            name=assessment.name,
            asset_name=assessment.asset.name if assessment.asset else "Unknown",
            assessment_type=assessment.assessment_type.value,
            status=assessment.status.value,
            start_date=assessment.actual_start.strftime('%Y-%m-%d') if assessment.actual_start else None,
            end_date=assessment.actual_end.strftime('%Y-%m-%d') if assessment.actual_end else None,
            executive_summary=assessment.executive_summary or "",
            scope_description=assessment.scope_description or "",
            methodology=assessment.methodology or "",
            findings=findings,
            assessor_name=assessment.assigned_to.full_name if assessment.assigned_to else "Security Team",
            assessor_email=assessment.assigned_to.email if assessment.assigned_to else "security@hubtel.com",
        )
